import docx
from docx.shared import Pt
from docx.shared import RGBColor as RGB
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os
import datetime


def set_cell_vertical_alignment(cell, align="center"):
    tc = cell._element
    tcPr = tc.get_or_add_tcPr()
    vAlign = OxmlElement('w:vAlign')
    vAlign.set(qn('w:val'), align)
    tcPr.append(vAlign)
    
def updateCells(filename, array, d):
    document = docx.Document(filename)
    table = document.tables[0]
    #if d.month % 3 == 1:
    col = 2
    #elif d.month % 3 == 2:
    #    col = 3
    #else:
    #    col = 4
    
    
            
    for i in range(2, 23):
        if not i == 4 and not i == 15 and not i == 21:
            table.rows[i].cells[col].text = (array[23] + ' (' if i == 22 else '') + array[i+3 if i<4 else i+2 if i<15 else i+1 if i<21 else i] + (')' if i == 22 else '')
            
        if i < 22: 
            if array[i + 21].endswith('Fail'):
                table.rows[i-3 if i<7 else i-2 if i<17 else i-1].cells[col].paragraphs[0].runs[0].font.color.rgb = RGB(255, 0, 0)
            else:
                table.rows[i].cells[col].paragraphs[0].runs[0].font.color.rgb = RGB(0, 0, 0)

    document.save(filename)




    
def createFile(filename, array, d):
    # Create a new document
    document = docx.Document()

    # Add a new section with Word 2016 compatibility settings
    section = document.sections[-1]
    section._sectPr = OxmlElement('w:sectPr')
    section._sectPr.set(qn('w:val'), 'continuous')
    section._sectPr.set(qn('w:docGrid'), '990')

    # Change the margins of the first section of the document
    for section in document.sections:
        #section = document.sections[0]
        section.left_margin = docx.shared.Inches(.75)
        section.right_margin = docx.shared.Inches(.75)
        section.top_margin = docx.shared.Inches(.5)
        section.bottom_margin = docx.shared.Inches(.5)

    # Set document styles
    style1 = document.styles['Normal']
    font1 = style1.font
    font1.name = 'Calibri'
    font1.size = docx.shared.Pt(12)
    style1.paragraph_format.space_after = docx.shared.Pt(6)

    new_style = document.styles.add_style('NewStyle', docx.enum.style.WD_STYLE_TYPE.PARAGRAPH)
    new_style.font.name = 'Calibri'
    new_style.font.bold = True
    new_style.font.size = docx.shared.Pt(28)
    new_style.paragraph_format.space_after = docx.shared.Pt(6)

    new_style2 = document.styles.add_style('NewStyle2', docx.enum.style.WD_STYLE_TYPE.PARAGRAPH)
    new_style2.font.name = 'Calibri'
    new_style2.font.bold = True
    new_style2.font.size = docx.shared.Pt(14)
    new_style2.paragraph_format.space_after = docx.shared.Pt(6)

    new_style3 = document.styles.add_style('NewStyle3', docx.enum.style.WD_STYLE_TYPE.PARAGRAPH)
    new_style3.font.name = 'Calibri'
    new_style3.font.bold = True
    new_style3.font.size = docx.shared.Pt(9)
    new_style3.paragraph_format.space_after = docx.shared.Pt(6)


    # Add heading to the document
    heading = document.add_paragraph('Data Quality Plan Quarterly Score Card')
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    heading.style = document.styles['NewStyle']
    heading.style.font.size = docx.shared.Pt(28)
    heading.style.font.bold = True
    heading.style.font.name = 'Calibri (Body)'

    # Add information about the quarter, agency, project, project type, and review date to the document
    quarter = array[24] + ' ' + array[25]
    agency = array[1]
    project = array[2]
    project_type = array[4]
    review_date = array[0]
    note = '*Does not contribute to grade'



    p1 = document.add_paragraph(f'Quarter: {quarter}', style = new_style2)
    p2 = document.add_paragraph(f'Agency: {agency}')
    p2.style = document.styles['NewStyle2']
    p3 = document.add_paragraph(f'Project: {project}')
    p3.style = document.styles['NewStyle2']
    p4 = document.add_paragraph(f'Project type: {project_type}')
    p4.style = document.styles['NewStyle2']
    p5 = document.add_paragraph(f'Review Date: {review_date}')
    p5.style = document.styles['NewStyle2']

    # Add the table to the document
    table = document.add_table(rows=23, cols=3, style='Light Shading')

    # Add the column headers to the table
    if array[24] == 'Q2':
        column_headers = ['', '', 'January - March']
    elif array[24] == 'Q3':
        column_headers = ['', '', 'April - June']
    elif array[24] == 'Q4':
        column_headers = ['', '', 'July - September']
    else:
        column_headers = ['', '', 'October - December']
    for i in range(3):
        table.cell(0, i).text = column_headers[i]

    # Add the rows to the table
    row_headers = [
        ('Timeliness', '', ''),
        ('', 'Timeliness', ''),
        ('', 'Enrollment Length (ES only)', ''),
        ('Completeness', '', ''),
        ('', 'Name', ''),
        ('', 'SSN', ''),
        ('', 'DOB', ''),
        ('', 'Race', ''),
        ('', 'Ethnicity', ''),
        ('', 'Gender', ''),
        ('', 'Vet Status', ''),
        ('', 'Exit Destination', ''),
        ('', 'Chronicity', ''),
        ('', 'Inactive Records (SO only)*', ''),
        ('Accuracy', '', ''),
        ('', 'Disabling Cond.', ''),
        ('', 'Income (Start)', ''),
        ('', 'Income (Annual)', ''),
        ('', 'Income (Exit)', ''),
        ('', 'Relationship to HoH', ''),
        ('', '', '', '', ''),
        ('Overall grade', '', '')
    ]

    for i in range(1, 23):
        for j in range(3):
            cell = table.cell(i, j)
            cell.text = row_headers[i-1][j]
            set_cell_vertical_alignment(cell, "center")
            cell.height = docx.shared.Cm(10.0)
            cell.margin_left = docx.shared.Inches(0.05)
            cell.margin_right = docx.shared.Inches(0.05)
            cell.margin_top = docx.shared.Inches(0.0)
            cell.margin_bottom = docx.shared.Inches(0.0)
            
            
    p6 = document.add_paragraph(f' {note}')
    p6.style = document.styles['NewStyle3']

    

    # Save the document
    document.save(filename)
    
    updateCells(filename, array, d)


def initialize(filename, array ):
    filename = os.getcwd() + "\\" + filename
    if len(array) > 0:
        d = datetime.datetime.strptime(array[0], "%m%d%Y")
        array[0]= d.strftime("%m/%d/%Y")
        #print(array[0])
    # Check if the file exists
    if not os.path.exists(filename):
        # Open the existing file
        createFile(filename, array, d)
    else:
        updateCells(filename, array, d)
        
